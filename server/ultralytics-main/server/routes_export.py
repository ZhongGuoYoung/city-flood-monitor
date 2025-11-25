# server/routes_export.py
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
from io import BytesIO
from pathlib import Path
from datetime import datetime
from urllib.parse import quote
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from starlette.responses import FileResponse

from .database import get_conn
from .utils.export_video import export_video_with_ticks

router = APIRouter(tags=["export"])


# ====== Pydantic 模型 ======

class CameraInfo(BaseModel):
    id: str = ""
    name: str = ""
    location: str = ""


class ExportRow(BaseModel):
    ts: int  # 毫秒
    level: int
    percent: float


class ExportRequest(BaseModel):
    camera: CameraInfo
    source_type: str  # "video"/"mp4" 或 "live"
    rows: List[ExportRow]


# ====== 时间格式化：mp4 用视频时间，监控用真实时间 ======

def format_time_label(ts_ms: int, source_type: str) -> str:
    st = (source_type or "").lower()
    if st in ("video", "mp4"):
        # 相对视频开始的时间
        total_sec = ts_ms // 1000
        h = total_sec // 3600
        m = (total_sec % 3600) // 60
        s = total_sec % 60
        return f"{h:02d}:{m:02d}:{s:02d}"
    else:
        # 实时监控：真实时间
        dt = datetime.fromtimestamp(ts_ms / 1000.0)
        return dt.strftime("%Y-%m-%d %H:%M:%S")


# 导出文件
# ====== 模板路径 ======

TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "video_result_template.docx"


def build_doc(req: ExportRequest) -> BytesIO:
    # 1. 打开模板
    if not TEMPLATE_PATH.exists():
        # 模板不存在，直接抛明白的错误
        raise HTTPException(status_code=500, detail=f"模板文件不存在：{TEMPLATE_PATH}")

    try:
        doc = Document(str(TEMPLATE_PATH))
    except PackageNotFoundError:
        raise HTTPException(status_code=500, detail=f"无法打开模板：{TEMPLATE_PATH}")

    # 2. 摄像头信息一行：摄像头信息：名称：xx，位置：yy，ID：zz
    cam = req.camera
    cam_info_text = f"摄像头信息：名称：{cam.name}，位置：{cam.location}，ID：{cam.id}"

    # 目前你的模板：0 = 标题，1 = 风险说明，2 = 空行，我们把第 2 段改成摄像头信息
    if len(doc.paragraphs) >= 3:
        p = doc.paragraphs[2]
        p.text = cam_info_text
    else:
        doc.add_paragraph(cam_info_text)

    # 3. 找第一张表作为结果表
    if not doc.tables:
        raise HTTPException(status_code=500, detail="模板中没有结果表，请检查模板格式")

    table = doc.tables[0]

    # 表头第一列改成「时间」
    header_cells = table.rows[0].cells
    header_cells[0].text = "时间"

    # 清掉除表头外的示例数据行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # 4. 写入数据行
    for r in req.rows:
        t_label = format_time_label(r.ts, req.source_type)
        risk_text = f"Level {r.level}"
        pct_text = f"{r.percent:.2f}%"

        cells = table.add_row().cells
        cells[0].text = t_label
        cells[1].text = risk_text
        cells[2].text = pct_text

    # 5. 存到内存流
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.post("/api/exportWord")
async def export_word(req: ExportRequest):
    # 防御：rows 为空直接 400，不要抛原始异常
    if not req.rows:
        raise HTTPException(status_code=400, detail="rows 不能为空")

    try:
        buf = build_doc(req)
    except HTTPException:
        # 已经是 HTTPException，直接往上抛
        raise
    except Exception as e:
        # 兜底：打印一下错误，返回 500
        print("[exportWord] unexpected error:", repr(e))
        raise HTTPException(status_code=500, detail=f"导出失败：{e}")

    # 构造下载文件名（可以有中文）
    base_name = (req.camera.name or "camera").strip() or "camera"
    file_name = f"{base_name}_识别结果.docx"

    # 用 RFC 5987 方式编码，保证 header 里都是 ASCII
    quoted_name = quote(file_name)  # 比如 "人民广场.docx" -> "%E4%BA%BA%E6%B0%91%E5%B9%BF%E5%9C%BA.docx"
    content_disposition = f"attachment; filename*=UTF-8''{quoted_name}"

    headers = {
        "Content-Disposition": content_disposition
    }

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )


# 导出视频
# ====== 导出带掩膜的视频 ======

BASE_DIR = Path(__file__).resolve().parent
RECORD_DIR = BASE_DIR / "records"
EXPORT_DIR = BASE_DIR / "exports"
EXPORT_DIR.mkdir(exist_ok=True)

@router.get("/api/history/{session_id}/exportVideo")
def export_history_video(session_id: int):
    """
    1. 从 detect_session 里查出视频相对路径
    2. 用 export_video_with_ticks 生成带掩膜的视频
    3. 返回 mp4 文件
    """
    conn = get_conn()
    try:
        with conn.cursor() as cur:
            sql = """
                SELECT record_path, camera_name
                FROM detect_session
                WHERE id = %s
            """
            cur.execute(sql, (session_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="session not found")
    finally:
        conn.close()

    rel_path = row["record_path"]
    src_path = RECORD_DIR / rel_path
    if not src_path.exists():
        raise HTTPException(status_code=404, detail="源视频不存在")

    out_path = EXPORT_DIR / (src_path.stem + "_mask.mp4")

    # 没生成过就生成一次，生成过就直接返回
    if not out_path.exists():
        try:
            export_video_with_ticks(session_id, str(src_path), str(out_path))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"导出失败: {e}")

    return FileResponse(
        path=out_path,
        media_type="video/mp4",
        filename=out_path.name,
    )
